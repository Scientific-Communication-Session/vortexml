"""
Vortex ML — Data Processor
Handles file uploads, Excel→CSV conversion, dataset analysis, and data preparation.
"""

import os
import json
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler, LabelEncoder
from werkzeug.utils import secure_filename
import torch
from torch.utils.data import DataLoader, TensorDataset

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


def safe_filename(name, default="upload.csv"):
    """Strip any path components / unsafe characters from an uploaded filename
    so it can never be joined into a path that escapes its target directory."""
    base = secure_filename(name or "")
    if not base:
        ext = os.path.splitext(name or "")[1]
        base = default if not ext else "upload" + ext
    return base


# Tabular-ish formats we accept for dataset upload. Everything that is not
# already a CSV is converted to one so the analyze/prepare pipeline downstream
# stays completely unchanged.
TABULAR_EXTS = (".csv", ".tsv", ".xlsx", ".xls", ".json", ".parquet",
                ".docx", ".html", ".htm")


def _to_dataframe(filepath, ext):
    """Load a non-CSV tabular file into a pandas DataFrame.

    Heavy/optional parsers are imported lazily inside each branch and guarded
    with a friendly ValueError so a missing dependency surfaces as a clean 400
    rather than an ImportError 500. Formats with no tabular content raise
    ValueError explaining there's no table to train on.
    """
    if ext == ".tsv":
        return pd.read_csv(filepath, sep="\t")

    if ext in (".xlsx", ".xls"):
        # .xls (legacy) needs xlrd; .xlsx uses openpyxl. Let pandas pick the
        # engine for .xls and fall back to a friendly message if it's missing.
        if ext == ".xlsx":
            return pd.read_excel(filepath, engine="openpyxl")
        try:
            return pd.read_excel(filepath)
        except ImportError:
            raise ValueError("Reading legacy .xls files needs the 'xlrd' "
                             "package. Re-save as .xlsx or .csv and try again.")

    if ext == ".json":
        # Accept an array of record objects (the common export shape) or a
        # column-oriented dict. Anything else has no table to train on.
        try:
            with open(filepath, "r", encoding="utf-8", errors="replace") as fh:
                data = json.load(fh)
        except (json.JSONDecodeError, ValueError) as e:
            raise ValueError(f"Could not parse this file as JSON: {e}")
        if isinstance(data, list):
            df = pd.json_normalize(data)
        elif isinstance(data, dict):
            # {"records": [...]} or {"data": [...]} wrappers, else treat the
            # dict itself as column->values.
            for key in ("records", "data", "rows"):
                if isinstance(data.get(key), list):
                    df = pd.json_normalize(data[key])
                    break
            else:
                df = pd.json_normalize(data)
        else:
            raise ValueError("This JSON file has no table to train on — "
                             "expected an array of objects (records).")
        if df.shape[1] == 0:
            raise ValueError("This JSON file has no columns to train on.")
        return df

    if ext == ".parquet":
        try:
            return pd.read_parquet(filepath)
        except ImportError:
            raise ValueError("Reading .parquet files needs the 'pyarrow' "
                             "package (pip install pyarrow).")

    if ext == ".docx":
        try:
            import docx  # python-docx
        except ImportError:
            raise ValueError("Reading tables from Word .docx files needs the "
                             "'python-docx' package (pip install python-docx).")
        document = docx.Document(filepath)
        if not document.tables:
            raise ValueError("This Word document has no table to train on. "
                             "Upload a .docx that contains a table, or use a "
                             "CSV/Excel file.")
        # Use the first table; first row is treated as the header.
        table = document.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            raise ValueError("The table in this Word document has no data rows "
                             "to train on.")
        return pd.DataFrame(rows[1:], columns=rows[0])

    if ext in (".html", ".htm"):
        # Parse <table> elements directly with BeautifulSoup (using the 'lxml'
        # parser). We avoid pandas.read_html here because in this pandas build it
        # pulls in the optional html5lib parser and raises a misleading
        # ImportError on table-less pages.
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ValueError("Reading tables from HTML needs the "
                             "'beautifulsoup4' and 'lxml' packages.")
        with open(filepath, "rb") as fh:
            soup = BeautifulSoup(fh.read(), "lxml")
        table = soup.find("table")
        if table is None:
            raise ValueError("This HTML file has no table to train on.")
        rows = []
        for tr in table.find_all("tr"):
            cells = tr.find_all(["th", "td"])
            if cells:
                rows.append([c.get_text(strip=True) for c in cells])
        if len(rows) < 2:
            raise ValueError("The table in this HTML file has no data rows "
                             "to train on.")
        # First row is the header; pad/truncate ragged rows to header width.
        header = rows[0]
        width = len(header)
        body = [(r + [""] * width)[:width] for r in rows[1:]]
        return pd.DataFrame(body, columns=header)

    # Should be unreachable — callers gate on TABULAR_EXTS.
    raise ValueError(f"Unsupported file type '{ext}'.")


def save_uploaded_file(file_storage):
    """Save an uploaded file, auto-converting tabular formats to CSV.

    CSV is saved as-is (existing flow, unchanged). Every other supported format
    (TSV, Excel, JSON, Parquet, Word/HTML tables) is loaded into a DataFrame and
    written out as CSV so the downstream analyze/prepare pipeline keeps working.
    """
    filename = safe_filename(file_storage.filename)
    filepath = os.path.join(UPLOAD_DIR, filename)
    file_storage.save(filepath)

    ext = os.path.splitext(filename)[1].lower()

    # CSV: keep exactly as before.
    if ext == ".csv":
        return filename, filepath

    if ext in TABULAR_EXTS:
        try:
            df = _to_dataframe(filepath, ext)
        except Exception:
            # Clean up the original upload on any failure so we don't leave a
            # stray non-CSV in the uploads dir, then re-raise for the caller.
            if os.path.exists(filepath):
                os.remove(filepath)
            raise
        csv_filename = os.path.splitext(filename)[0] + ".csv"
        csv_path = os.path.join(UPLOAD_DIR, csv_filename)
        df.to_csv(csv_path, index=False)
        # Remove the original non-CSV upload (unless it WAS the csv target name).
        if os.path.exists(filepath) and filepath != csv_path:
            os.remove(filepath)
        return csv_filename, csv_path

    return filename, filepath


def _read_csv(path):
    """Read a user-uploaded CSV robustly.

    Real-world CSVs (especially Excel exports) are frequently Windows-1252 /
    Latin-1 rather than UTF-8, so a strict utf-8 read raises UnicodeDecodeError
    on a stray byte. Fall back to latin-1, which maps all 256 byte values and
    never fails to decode. Empty/columnless or otherwise unparseable files are
    re-raised as ValueError with a friendly message so callers can return a 400
    instead of a 500 with a raw stack trace.
    """
    try:
        return pd.read_csv(path)
    except UnicodeDecodeError:
        return pd.read_csv(path, encoding="latin-1")
    except (pd.errors.EmptyDataError, pd.errors.ParserError) as e:
        raise ValueError(f"Could not read this file as a CSV: {e}")


def analyze_dataset(csv_path):
    """Analyze a CSV and return metadata for the Dataset Designer UI."""
    df = _read_csv(csv_path)

    columns = []
    for col in df.columns:
        col_info = {
            "name": col,
            "dtype": str(df[col].dtype),
            "non_null": int(df[col].notna().sum()),
            "null_count": int(df[col].isna().sum()),
            "unique": int(df[col].nunique()),
            "is_numeric": pd.api.types.is_numeric_dtype(df[col]),
        }
        if col_info["is_numeric"]:
            col_info["min"] = float(df[col].min()) if not df[col].isna().all() else None
            col_info["max"] = float(df[col].max()) if not df[col].isna().all() else None
            col_info["mean"] = float(df[col].mean()) if not df[col].isna().all() else None
            col_info["std"] = float(df[col].std()) if not df[col].isna().all() else None
        else:
            top_values = df[col].value_counts().head(5).to_dict()
            col_info["top_values"] = {str(k): int(v) for k, v in top_values.items()}
        columns.append(col_info)

    # Preview rows (first 10)
    preview = df.head(10).fillna("").to_dict(orient="records")

    return {
        "rows": len(df),
        "cols": len(df.columns),
        "columns": columns,
        "preview": preview,
    }


def _split(X, y, test_size, stratify):
    """train_test_split that uses stratification when possible.

    Stratifying keeps the class balance identical across train/val/test, which
    matters for imbalanced classification. It requires >= 2 samples per class in
    the split, so we fall back to an unstratified split if a class is too rare
    (otherwise sklearn raises and training never starts).
    """
    if stratify is not None:
        try:
            return train_test_split(X, y, test_size=test_size,
                                    random_state=42, stratify=stratify)
        except ValueError:
            pass
    return train_test_split(X, y, test_size=test_size, random_state=42)


def dataset_health(csv_path, feature_cols=None, target_col=None):
    """Surface data-quality risks before training.

    Checks (column-level always; target-aware when a target is given):
      - tiny dataset, high-null columns, constant / near-constant features,
        ID-like and high-cardinality categoricals,
      - class imbalance (classification target),
      - likely target leakage (a feature that perfectly determines the target).

    Returns {"rows", "warnings": [{level, code, title, detail, columns}], "ok"}.
    Everything is best-effort and read-only.
    """
    df = _read_csv(csv_path)
    n = len(df)
    feature_cols = [c for c in (feature_cols or []) if c in df.columns]
    warnings = []

    def add(level, code, title, detail, columns=None):
        warnings.append({"level": level, "code": code, "title": title,
                         "detail": detail, "columns": columns or []})

    if n < 200:
        add("warning", "tiny_dataset", "Very small dataset",
            f"Only {n} rows. Models overfit easily here — prefer a simple "
            f"architecture (MLP) and keep early stopping on.")

    cols_to_check = feature_cols or [c for c in df.columns if c != target_col]
    for col in cols_to_check:
        s = df[col]
        nn = int(s.notna().sum())
        null_frac = (1 - nn / n) if n else 0
        nunique = int(s.nunique(dropna=True))

        if null_frac > 0.3:
            add("warning", "high_null", f"'{col}' is mostly empty",
                f"{round(null_frac * 100)}% of '{col}' is missing; it will be "
                f"filled in, which can mislead the model.", [col])

        if nunique <= 1:
            add("warning", "constant", f"'{col}' is constant",
                f"'{col}' has a single value — no signal. Consider removing it.", [col])
        elif nn:
            top_frac = float(s.value_counts(normalize=True, dropna=True).iloc[0])
            if top_frac > 0.99:
                add("info", "near_constant", f"'{col}' is nearly constant",
                    f"{round(top_frac * 100)}% of '{col}' is one value — little signal.", [col])

        if not pd.api.types.is_numeric_dtype(s):
            if nunique == n and n > 0:
                add("warning", "id_like", f"'{col}' looks like an identifier",
                    f"'{col}' is unique on every row — likely an ID that can't "
                    f"generalize. Consider excluding it.", [col])
            elif nunique > 50 and nn and (nunique / nn) > 0.5:
                add("warning", "high_cardinality", f"'{col}' is high-cardinality",
                    f"'{col}' has {nunique} distinct text values; encoding it adds "
                    f"noise. Consider grouping rare values or dropping it.", [col])

    if target_col and target_col in df.columns:
        tgt = df[target_col]
        t_unique = int(tgt.nunique(dropna=True))
        task = "regression" if (pd.api.types.is_numeric_dtype(tgt) and t_unique > 10) else "classification"

        if task == "classification":
            counts = tgt.value_counts(normalize=True, dropna=True)
            if len(counts) and counts.min() < 0.1 and len(counts) <= 50:
                add("warning", "class_imbalance", "Imbalanced target classes",
                    f"Class '{counts.idxmin()}' is only {round(float(counts.min()) * 100, 1)}% "
                    f"of rows (largest is {round(float(counts.max()) * 100)}%). Accuracy can "
                    f"be misleading — watch the per-class results.", [target_col])
            leaky = []
            for col in feature_cols:
                if col == target_col:
                    continue
                fu = int(df[col].nunique(dropna=True))
                if 1 < fu <= 50:
                    grp = df.groupby(col, dropna=True)[target_col].nunique()
                    if len(grp) and (grp <= 1).all():
                        leaky.append(col)
            if leaky:
                add("warning", "leakage", "Possible target leakage",
                    f"{', '.join(leaky)} perfectly predict '{target_col}' in this data. "
                    f"If they wouldn't be known before the outcome, training on them "
                    f"inflates accuracy unrealistically.", leaky)
        else:
            leaky = []
            for col in feature_cols:
                if col == target_col or not pd.api.types.is_numeric_dtype(df[col]):
                    continue
                try:
                    pair = df[[col, target_col]].dropna()
                    if len(pair) > 2 and abs(float(pair.corr().iloc[0, 1])) > 0.98:
                        leaky.append(col)
                except Exception:
                    pass
            if leaky:
                add("warning", "leakage", "Possible target leakage",
                    f"{', '.join(leaky)} are almost perfectly correlated with "
                    f"'{target_col}' — likely leakage.", leaky)

    return {
        "rows": n,
        "warnings": warnings,
        "ok": not any(w["level"] == "warning" for w in warnings),
    }


def prepare_dataset(csv_path, feature_cols, target_col, batch_size=32, test_size=0.2, val_size=0.1):
    """
    Prepare a dataset for training.
    Returns train_loader, val_loader, test_loader, input_dim, output_dim, task_type.
    """
    df = _read_csv(csv_path)

    # Determine task type
    target_series = df[target_col]
    if pd.api.types.is_numeric_dtype(target_series) and target_series.nunique() > 10:
        task_type = "regression"
    else:
        task_type = "classification"

    # Encode categorical feature columns. `features_meta` records exactly how
    # each column was transformed so the same mapping can be replayed at
    # inference time (see apply_preprocess).
    label_encoders = {}
    features_meta = {}
    for col in feature_cols:
        if not pd.api.types.is_numeric_dtype(df[col]):
            le = LabelEncoder()
            df[col] = le.fit_transform(df[col].astype(str).fillna("__MISSING__"))
            label_encoders[col] = le
            features_meta[col] = {"kind": "categorical",
                                  "classes": [str(c) for c in le.classes_]}
        else:
            median = float(df[col].median()) if not df[col].isna().all() else 0.0
            df[col] = df[col].fillna(median)
            features_meta[col] = {"kind": "numeric", "median": median}

    # Prepare features
    X = df[feature_cols].values.astype(np.float32)

    # Prepare target
    if task_type == "classification":
        le_target = LabelEncoder()
        y = le_target.fit_transform(target_series.astype(str).fillna("__MISSING__"))
        output_dim = len(le_target.classes_)
        y = y.astype(np.int64)
        target_classes = [str(c) for c in le_target.classes_]
    else:
        y = target_series.fillna(target_series.median()).values.astype(np.float32)
        output_dim = 1
        target_classes = None

    # Split FIRST, then fit the scaler on the training split only. Fitting on
    # the full matrix leaks validation/test statistics (mean, std) into the
    # training distribution and inflates reported metrics.
    strat = y if task_type == "classification" else None
    X_train, X_test, y_train, y_test = _split(X, y, test_size, strat)
    strat_train = y_train if task_type == "classification" else None
    X_train, X_val, y_train, y_val = _split(X_train, y_train, val_size, strat_train)

    # Scale features — fit on train, apply the same transform to val/test.
    scaler = StandardScaler()
    X_train = scaler.fit_transform(X_train)
    X_val = scaler.transform(X_val)
    X_test = scaler.transform(X_test)

    # Create DataLoaders
    def make_loader(X_arr, y_arr, shuffle, drop_last=False):
        X_t = torch.tensor(X_arr)
        y_t = torch.tensor(y_arr)
        ds = TensorDataset(X_t, y_t)
        return DataLoader(ds, batch_size=batch_size, shuffle=shuffle, drop_last=drop_last)

    # BatchNorm layers (DNN/ResNet) raise "Expected more than 1 value per
    # channel" on a batch of exactly 1 sample in training mode. If the train
    # split would leave a single-sample final batch, drop it — but only when
    # there is more than one batch, so we never empty the loader on small data.
    # Validation/test run under model.eval(), where BatchNorm is fine with a
    # size-1 batch, so only the training loader needs this.
    train_drop_last = (len(X_train) % batch_size == 1) and (len(X_train) > batch_size)
    train_loader = make_loader(X_train, y_train, shuffle=True, drop_last=train_drop_last)
    val_loader = make_loader(X_val, y_val, shuffle=False)
    test_loader = make_loader(X_test, y_test, shuffle=False)

    input_dim = len(feature_cols)

    # Everything needed to reproduce this exact preprocessing on new rows.
    preprocess = {
        "feature_cols": list(feature_cols),
        "target_col": target_col,
        "task_type": task_type,
        "input_dim": input_dim,
        "output_dim": output_dim,
        "scaler": {"mean": scaler.mean_.tolist(), "scale": scaler.scale_.tolist()},
        "features": features_meta,
        "target_classes": target_classes,
    }

    return {
        "train_loader": train_loader,
        "val_loader": val_loader,
        "test_loader": test_loader,
        "input_dim": input_dim,
        "output_dim": output_dim,
        "task_type": task_type,
        "train_size": len(X_train),
        "val_size": len(X_val),
        "test_size": len(X_test),
        "preprocess": preprocess,
    }


# ─────────────────────────────────────────────────────────
# Inference — replay the training preprocessing on new rows
# ─────────────────────────────────────────────────────────
def preprocess_sidecar_path(weights_path):
    """The JSON sidecar path for a given .pt weights file (same base name)."""
    base = weights_path[:-3] if weights_path.endswith(".pt") else weights_path
    return base + ".preprocess.json"


def save_preprocess(weights_path, preprocess):
    """Persist preprocessing metadata next to the weights file."""
    with open(preprocess_sidecar_path(weights_path), "w", encoding="utf-8") as f:
        json.dump(preprocess, f)


def load_preprocess(weights_path):
    """Load preprocessing metadata for a weights file, or None if absent."""
    path = preprocess_sidecar_path(weights_path)
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def apply_preprocess(preprocess, rows):
    """Turn a list of {col: value} dicts into a scaled float32 feature matrix,
    replaying the exact transforms recorded at training time.

    Unknown categorical values fall back to the '__MISSING__' bucket (or 0);
    missing/non-numeric numeric values fall back to the stored training median.
    """
    feature_cols = preprocess["feature_cols"]
    feats = preprocess["features"]
    mean = np.array(preprocess["scaler"]["mean"], dtype=np.float32)
    scale = np.array(preprocess["scaler"]["scale"], dtype=np.float32)

    matrix = []
    for row in rows:
        encoded = []
        for col in feature_cols:
            meta = feats.get(col, {"kind": "numeric", "median": 0.0})
            raw = row.get(col)
            if meta["kind"] == "categorical":
                classes = meta["classes"]
                key = "__MISSING__" if raw is None else str(raw)
                if key in classes:
                    encoded.append(float(classes.index(key)))
                elif "__MISSING__" in classes:
                    encoded.append(float(classes.index("__MISSING__")))
                else:
                    encoded.append(0.0)
            else:
                try:
                    val = float(raw)
                    if np.isnan(val):
                        val = meta["median"]
                except (TypeError, ValueError):
                    val = meta["median"]
                encoded.append(val)
        matrix.append(encoded)

    X = np.array(matrix, dtype=np.float32)
    return (X - mean) / scale
